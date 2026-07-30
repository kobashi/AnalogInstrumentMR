"""Build the image-rich AI review package for current 3D instrument models."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "AnalogInstrumentMR_3Dモデル_AIレビュー資料.docx"
RUNTIME_IMAGES = ROOT / "docs" / "review_assets" / "runtime_greyboxes"
FONT = "YuGothic"

# Selected style system: compact_reference_guide.
# Named overrides: YuGothic for Japanese coverage; 9 pt compact data tables;
# 24 pt editorial-cover title; full-page model-card page breaks.
INK = "14283A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64727E"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
WHITE = "FFFFFF"
AMBER = "F39B32"
GREEN = "40D98A"
RED = "D95252"

THEMES = [
    ("OrbitalAnalog", "ORBITAL ANALOG", "orbital-analog"),
    ("ForgeBrass", "FORGE BRASS", "forge-brass"),
    ("KineticSafety", "KINETIC SAFETY", "kinetic-safety"),
]

MODELS = [
    {
        "key": "RoundMeter",
        "image": "MeterRound",
        "name": "円形アナログメーター",
        "display": "ROUND METER",
        "type_id": "meter.round",
        "role": "読取専用。0–1の値を針で示し、操作モード中は微動する。",
        "envelope": "0.140 × 0.140 × 0.064 m",
        "motion": "needle ±55°／微動アニメーション",
        "interaction": "操作不可（Connect出力候補）",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "1–3 mから円形メーターとして即時に読めるシルエットと、近距離での目盛り密度。",
            "ベゼル、盤面、針の奥行き階層がテーマごとに明確か。針とhubが細すぎないか。",
            "3テーマが同一機能を保ちながら、形状言語として十分に差別化されているか。",
        ],
    },
    {
        "key": "Lever",
        "image": "Lever",
        "name": "5段階レバー",
        "display": "LEVER",
        "type_id": "control.lever",
        "role": "5段階の操作入力。ビーム＋Trigger、またはGrip＋controller motion。",
        "envelope": "0.180 × 0.256 × 0.100 m",
        "motion": "local X、±24°、5 detents、初期offset −24°",
        "interaction": "Ray／Grip motion",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "pivotが基部から浮かず、全可動域でhandleがhousingへめり込まない機構表現。",
            "面に垂直な操作面、手で握る位置、レバー長がQuest controller操作に自然か。",
            "5段階のdetentを形状・目盛り・guardで予告できるか。",
        ],
    },
    {
        "key": "ToggleSwitch",
        "image": "Toggle",
        "name": "トグルスイッチ",
        "display": "TOGGLE",
        "type_id": "control.toggle",
        "role": "ON／OFF入力。短い操作とクリック感を想定。",
        "envelope": "0.120 × 0.170 × 0.064 m",
        "motion": "local X、±28°、初期offset −28°",
        "interaction": "Ray＋Trigger",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "回転起点がbaseへ接続して見え、初期角度と反転時のどちらでも埋没しないか。",
            "細いshaftでも遠距離からON／OFF方向が読み取れるか。",
            "collar・guardが操作を妨げず、押しボタンやレバーと取り違えないか。",
        ],
    },
    {
        "key": "RotaryKnob",
        "image": "Rotary",
        "name": "ロータリーノブ",
        "display": "ROTARY KNOB",
        "type_id": "control.rotary",
        "role": "連続値または段階値の入力。",
        "envelope": "0.150 × 0.150 × 0.102 m",
        "motion": "local forward軸、連続360°",
        "interaction": "Ray＋Trigger",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "knurl風の外周、index、collarが小型表示でも操作方向を伝えるか。",
            "controller beamで選択する面積と、実物らしい指掛かりのバランス。",
            "回転状態を盤面側の目盛りとノブ側indexで追跡できるか。",
        ],
    },
    {
        "key": "PushButton",
        "image": "Button",
        "name": "押しボタン",
        "display": "PUSH BUTTON",
        "type_id": "control.button",
        "role": "Momentary入力。controller接触とビーム＋Triggerの両方に反応。",
        "envelope": "0.130 × 0.130 × 0.090 m",
        "motion": "面法線方向、travel 0.014 m",
        "interaction": "Contact press／Ray＋Trigger",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "capとguardの高低差が押下可能性を示し、接触操作を妨げないか。",
            "14 mmの沈み込みが視覚的に十分で、解放状態と混同しないか。",
            "ランプとの識別が色だけでなくflat cap／domeの形状でも成立するか。",
        ],
    },
    {
        "key": "IndicatorLamp",
        "image": "Lamp",
        "name": "単色状態ランプ",
        "display": "STATUS LAMP",
        "type_id": "indicator.lamp",
        "role": "ON／OFF表示。Connect出力として発光・pulseに利用。",
        "envelope": "0.140 × 0.110 × 0.082 m",
        "motion": "変形なし／単色発光pulse",
        "interaction": "操作不可（Connect出力候補）",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "dome lensとcage／ringの関係が発光器具として自然か。",
            "OFFでも輪郭を失わず、ONではrealtime lightなしで十分に目立つか。",
            "押しボタンと同じaccent色を使っても、シルエットで明確に区別できるか。",
        ],
    },
    {
        "key": "StatusIndicator",
        "image": "StatusIndicator",
        "name": "多段階状態LED",
        "display": "STATUS INDICATOR",
        "type_id": "indicator.status",
        "role": "OFF／SAFE／WARN／DANGERの4状態を表示。",
        "envelope": "0.180 × 0.120 × 0.075 m",
        "motion": "変形なし／黒・緑・橙・赤の状態発光",
        "interaction": "操作不可（Connect出力候補）",
        "asset": "runtime primitive greybox（Blender原本・LODは未制作）",
        "focus": [
            "4状態を色だけに依存せず、形・区画・ラベルで補強する余地。",
            "テーマ別lens形状（横長／宝石状／警告バー）の差が十分か。",
            "OFF時の存在感とDANGER時の強さを、発光量を増やしすぎず両立できるか。",
        ],
    },
    {
        "key": "ThrottleLever",
        "image": "Throttle",
        "name": "スロットルレバー",
        "display": "THROTTLE LEVER",
        "type_id": "control.throttle",
        "role": "CUTOFF～FULLの6段階エンジン出力入力。",
        "envelope": "0.240 × 0.340 × 0.160 m",
        "motion": "local X、片側70° sweep（実装amplitude ±35°＋offset）",
        "interaction": "Ray／Grip＋arc motion",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "航空機・船舶のengine quadrantらしいpalm grip、fork arm、軌道表現。",
            "視覚上のgrip位置とcontrollerが反応する位置が一致して見えるか。",
            "6段階の目盛り・端点・CUTOFFを、過度な細部なしで伝えられるか。",
        ],
    },
    {
        "key": "PowerSlider",
        "image": "PowerSlider",
        "name": "パワースライダー",
        "display": "POWER SLIDER",
        "type_id": "control.power_slider",
        "role": "OFF～MAXを10%刻みで操作する11段階入力。",
        "envelope": "0.160 × 0.340 × 0.094 m",
        "motion": "local up、travel 0.180 m、11 detents",
        "interaction": "Ray／Grip＋linear motion",
        "asset": "Blender原本・FBX・Unity prefab 実装済み",
        "focus": [
            "slot、rail、carriageの軸が一致し、移動中もhandleがframeから外れて見えないか。",
            "10%刻みを必要な距離から読み取れる目盛り／節度表現。",
            "握れる横長handleと、panel内で場所を取りすぎない外形のバランス。",
        ],
    },
    {
        "key": "WindowMeter",
        "image": "WindowMeter",
        "name": "窓枠サイズ大型メーター",
        "display": "WINDOW METER",
        "type_id": "meter.window",
        "role": "宇宙船内装のprimary silhouetteとなる大型読取計器。",
        "envelope": "1.200 × 0.750 × 0.190 m",
        "motion": "大型needle ±55°／微動アニメーション",
        "interaction": "操作不可（Connect出力候補）",
        "asset": "runtime primitive greybox（実モデル・LODは未制作）",
        "focus": [
            "小型meterの単純拡大に見えず、埋込筐体・深いdial・支持構造が成立するか。",
            "壁・床・天井を区別しない任意面配置でも、上下に依存しすぎない構成か。",
            "1–3 mの視認性と、4 renderer以下のQuest予算を両立する造形案。",
        ],
    },
    {
        "key": "WindowPanel",
        "image": "WindowPanel",
        "name": "窓枠サイズ宇宙船パネル",
        "display": "WINDOW PANEL",
        "type_id": "panel.window",
        "role": "隔壁・床・天井に配置する大型status panel。",
        "envelope": "1.600 × 0.900 × 0.220 m",
        "motion": "status vane ±42°／微動アニメーション",
        "interaction": "操作不可（Connect出力候補）",
        "asset": "runtime primitive greybox（実モデル・LODは未制作）",
        "focus": [
            "広いnegative space、左右rail、vaneが宇宙船設備として説得力を持つか。",
            "どの面に置いても成立し、暗黙の『上』へ依存しすぎないデザインか。",
            "将来モニター表示を組み込む余地と、現行status vaneの役割の整理。",
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(
    table, widths, indent=120, cell_margin=80, side_margin=120
):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (
                ("top", cell_margin),
                ("bottom", cell_margin),
                ("start", side_margin),
                ("end", side_margin),
            ):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=11, color=INK, bold=False):
    for run in paragraph.runs:
        set_font(run, size=size, color=color, bold=bold)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr_text, fld_sep, text, fld_end))
    set_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "AnalogInstrumentMR  |  3D MODEL REVIEW PACKAGE"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(hp, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lead = fp.add_run("CONFIDENTIAL · DEVELOPMENT REVIEW  |  ")
    set_font(lead, size=8.5, color=MUTED)
    add_page_field(fp)


def add_kicker(doc, text, color=AMBER, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=9, color=color, bold=True)
    run.font.all_caps = True
    return p


def add_title(doc, text, size=24, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=size, color=INK, bold=True)
    return p


def add_subtitle(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, size=size, color=MUTED)
    return p


def add_shaded_note(
    doc, title, body, fill=CALL_OUT, font_size=10, after=8
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    r1 = p.add_run(f"{title}  ")
    set_font(r1, size=font_size, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    set_font(r2, size=font_size, color=INK)


def add_bullets(doc, items, size=10.5, after=4, line_spacing=1.25):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line_spacing
        marker = p.add_run("• ")
        set_font(marker, size=size, color=BLUE, bold=True)
        run = p.add_run(item)
        set_font(run, size=size, color=INK)


def image_path(model, theme_folder):
    if model["key"] in ("StatusIndicator", "WindowMeter", "WindowPanel"):
        return RUNTIME_IMAGES / f"Preview_{model['image']}_{theme_folder}.png"
    return (
        ROOT
        / "ArtSource"
        / "Blender"
        / theme_folder
        / f"Preview_{model['image']}_{theme_folder}.png"
    )


def set_alt_text(inline_shape, text):
    inline_shape._inline.docPr.set("descr", text)
    inline_shape._inline.docPr.set("title", text)


def add_image(paragraph, path, width, alt_text):
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    set_alt_text(shape, alt_text)
    return shape


def add_theme_image_table(doc, model, image_width=1.48):
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_geometry(
        table, [3120, 3120, 3120], cell_margin=35, side_margin=60
    )
    set_repeat_table_header(table.rows[0])
    for idx, (folder, display, theme_id) in enumerate(THEMES):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(display)
        set_font(run, size=8.2, bold=True, color=DARK_BLUE)

        image_cell = table.cell(1, idx)
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p2 = image_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        path = image_path(model, folder)
        if not path.exists():
            raise FileNotFoundError(path)
        add_image(
            p2,
            path,
            image_width,
            f"{model['name']}の{display}テーマ現行モデル",
        )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(4)
    label = (
        "図：runtimeコードから再現した現行greybox（未制作のBlenderモデルではない）"
        if model["key"] in ("StatusIndicator", "WindowMeter", "WindowPanel")
        else "図：実装済みBlender原本の現行preview render"
    )
    run = caption.add_run(label)
    set_font(run, size=7.8, color=MUTED, italic=True)


def add_spec_table(doc, model):
    rows = [
        ("ID／用途", f"{model['type_id']}　|　{model['role']}"),
        ("外形／可動", f"{model['envelope']}　|　{model['motion']}"),
        (
            "操作／配置",
            f"{model['interaction']}　|　Plane／Volumeの任意面（上下の区別なし）",
        ),
        ("現行資産", model["asset"]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_table_geometry(
        table, [1700, 7660], cell_margin=35, side_margin=80
    )
    for row, (label, detail) in zip(table.rows, rows):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], PALE_GRAY)
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_font(lr, size=8.2, bold=True, color=DARK_BLUE)
        dp = row.cells[1].paragraphs[0]
        dp.paragraph_format.space_after = Pt(0)
        dr = dp.add_run(detail)
        set_font(dr, size=8.2, color=INK)


def report_triangles(model):
    if model["key"] in ("StatusIndicator", "WindowMeter", "WindowPanel"):
        return "runtime primitive greybox／renderer budget 2・4・4"
    values = []
    for folder, display, theme_id in THEMES:
        report = (
            ROOT
            / "ArtSource"
            / "Blender"
            / folder
            / f"BL_{model['image']}_{folder}.report.json"
        )
        data = json.loads(report.read_text(encoding="utf-8"))
        values.append(f"{display}: {data['triangles']:,} tris")
    return "／".join(values)


def add_model_page(doc, model, index):
    kicker = add_kicker(
        doc, f"MODEL {index:02d} / {len(MODELS):02d} · {model['type_id']}"
    )
    kicker.paragraph_format.page_break_before = True
    add_title(doc, f"{model['display']}  |  {model['name']}", size=17)
    add_theme_image_table(doc, model)
    add_spec_table(doc, model)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(3)
    r1 = note.add_run("実装予算メモ：")
    set_font(r1, size=7.8, bold=True, color=MUTED)
    r2 = note.add_run(report_triangles(model))
    set_font(r2, size=7.8, color=MUTED)
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(3)
    h.paragraph_format.space_after = Pt(2)
    hr = h.add_run("このモデルで確認してほしい点")
    set_font(hr, size=10.5, bold=True, color=BLUE)
    add_bullets(doc, model["focus"], size=8.4)
    add_shaded_note(
        doc,
        "回答フォーマット",
        "総合1–5点／良い点／最優先の変更3件／形状案／材質案／操作・可動干渉／他モデルとの統一性",
        fill=CALL_OUT,
        font_size=8.2,
        after=2,
    )


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(54)
    add_kicker(doc, "CURRENT IMPLEMENTATION · EXTERNAL AI REVIEW", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_title(doc, "AnalogInstrumentMR", size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_title(doc, "3Dモデル AIレビュー資料", size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_subtitle(
        doc,
        "実装済み計器・操作部品 11種類 × 3テーマ\n"
        "形状／材質／可動部／MR操作性のブラッシュアップ用",
        size=12.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    cover_models = [MODELS[0], MODELS[7], MODELS[10]]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120], indent=120)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, (model, theme) in enumerate(zip(cover_models, THEMES)):
        cell = table.cell(0, idx)
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        add_image(
            p2,
            image_path(model, theme[0]),
            1.88,
            f"表紙：{model['name']} {theme[1]}",
        )
        label = cell.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.paragraph_format.space_after = Pt(0)
        run = label.add_run(model["display"])
        set_font(run, size=8.5, bold=True, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(24)
    meta.paragraph_format.space_after = Pt(0)
    r = meta.add_run(
        f"Snapshot: {date.today().isoformat()}  |  Repository base: a11169f  |  Unity / Meta Quest"
    )
    set_font(r, size=9, color=MUTED)


def add_review_brief(doc):
    kicker = add_kicker(doc, "REVIEW BRIEF")
    kicker.paragraph_format.page_break_before = True
    add_title(doc, "レビューの目的と前提", size=22)
    add_shaded_note(
        doc,
        "目的",
        "現行の機能・寸法・可動契約を維持しつつ、SFアニメ／映画の戦艦ブリッジやコックピットのpropとして、触りたくなる造形へブラッシュアップする。",
    )
    doc.add_heading("現在の実装範囲", level=2)
    add_bullets(
        doc,
        [
            "11種類を配置可能。8種類は3テーマのBlender原本・FBX・Unity prefabを実装済み。",
            "多段階状態LED、大型メーター、大型パネルの3種類はruntime primitive greybox。今回の画像は実装コードを再現したreview render。",
            "全オブジェクトはPlane／Volumeの任意面へ配置でき、このアプリでは上下を区別しない。",
            "操作モードでは両手controllerを使用。Ray＋Trigger、Grip＋motion、button接触を部品ごとに使い分ける。",
            "Connectモードでは操作部品を入力、メーター・LED・パネルを出力として連動させる構想。",
        ],
    )
    doc.add_heading("レビュー時に固定する条件", level=2)
    add_bullets(
        doc,
        [
            "共通envelope、mount面、pivot、可動軸、InteractionColliderはテーマ間で変更しない。",
            "小型部品は5,000 tris以下、4 renderer以下、shared material 2以下。大型はLOD前提。",
            "計器ごとのrealtime lightは使用しない。透明は必要最小限、発光はshared material＋parameter。",
            "特定作品の固有ロゴ・特徴的形状・panel配置は複製せず、独自の形状言語として提案する。",
        ],
    )
    criteria_kicker = add_kicker(doc, "REVIEW CRITERIA")
    criteria_kicker.paragraph_format.page_break_before = True
    add_title(doc, "8つの評価軸", size=22)
    add_subtitle(
        doc,
        "各モデルは下記の質問で1–5点評価し、形状・材質・操作性の提案を分離して回答する。",
        size=10.5,
    )
    axes = [
        ("1", "Silhouette", "1–3 mで用途を判別できるか"),
        ("2", "Mechanical plausibility", "mount、pivot、支持、可動clearanceが自然か"),
        ("3", "Affordance", "握る・押す・回す方向を説明なしで理解できるか"),
        ("4", "MR ergonomics", "controller ray／grip位置と見た目が一致するか"),
        ("5", "Theme language", "3テーマの差とfamily内統一が成立するか"),
        ("6", "Material readability", "opaque／metal／emissiveの役割が明確か"),
        ("7", "Quest efficiency", "少ないmesh・materialでも密度感を作れるか"),
        ("8", "Originality", "既存IPを模倣せずSF propとして魅力があるか"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(
        table, [720, 2500, 6140], cell_margin=45, side_margin=80
    )
    headers = ("#", "評価軸", "判定質問")
    for idx, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], PALE_BLUE)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(text)
        set_font(rr, size=9, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for n, label, question in axes:
        cells = table.add_row().cells
        for idx, text in enumerate((n, label, question)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(text)
            set_font(rr, size=8.8, color=INK, bold=(idx == 1))
    add_shaded_note(
        doc,
        "判定の目安",
        "3点＝機能は伝わるが造形上の改善余地あり。4点＝現行制約内で十分に魅力的。5点＝他のモデルの基準として展開できる。",
        fill=PALE_BLUE,
    )


def add_theme_page(doc):
    kicker = add_kicker(doc, "ART DIRECTION")
    kicker.paragraph_format.page_break_before = True
    add_title(doc, "3テーマの形状言語", size=22)
    add_subtitle(
        doc,
        "同一の機能・寸法・pivotを保ち、silhouette、depth、palette、surface detailで差別化する。",
        size=10.5,
    )
    theme_models = [MODELS[0], MODELS[7], MODELS[5]]
    descriptions = {
        "OrbitalAnalog": (
            "薄い黒panel、深い円形dial、細いshaft、warm amber／原色indicator。"
            "細線・arc・tickはatlas中心。"
        ),
        "ForgeBrass": (
            "鋳鉄の暗い母材、段付きrim、aged brass／copper accent、rivets、球grip。"
        ),
        "KineticSafety": (
            "graphite shroud、面取りguard、太い支持材、orange／yellow safety accent。"
        ),
    }
    for theme_index, (folder, display, theme_id) in enumerate(THEMES):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(2)
        hr = h.add_run(f"{theme_index + 1}. {display}")
        set_font(hr, size=10.5, bold=True, color=BLUE)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        rr = p.add_run(descriptions[folder])
        set_font(rr, size=8.4, color=INK)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        set_repeat_table_header(table.rows[0])
        set_table_geometry(
            table, [3120, 3120, 3120], cell_margin=25, side_margin=40
        )
        for idx, model in enumerate(theme_models):
            cell = table.cell(0, idx)
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_after = Pt(0)
            add_image(
                pp,
                image_path(model, folder),
                1.05,
                f"{display}テーマの{model['name']}",
            )
    add_shaded_note(
        doc,
        "横断レビュー",
        "各テーマ内でmeter・control・indicatorが同じ製造体系に見えるか、また色を失ってもテーマを識別できるかを評価する。",
        font_size=8.2,
        after=0,
    )


def add_summary_page(doc):
    kicker = add_kicker(doc, "IMPLEMENTATION MAP")
    kicker.paragraph_format.page_break_before = True
    add_title(doc, "モデル一覧と制作状態", size=22)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_geometry(
        table,
        [500, 1800, 1940, 2480, 2640],
        cell_margin=24,
        side_margin=45,
    )
    headers = ("#", "Type ID", "名称", "状態", "主な可動／状態")
    for idx, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], PALE_BLUE)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(text)
        set_font(rr, size=7.8, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for idx, model in enumerate(MODELS, 1):
        status = (
            "Runtime greybox"
            if model["key"] in ("StatusIndicator", "WindowMeter", "WindowPanel")
            else "3-theme model"
        )
        values = (str(idx), model["type_id"], model["name"], status, model["motion"])
        cells = table.add_row().cells
        for col, text in enumerate(values):
            p = cells[col].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(text)
            set_font(rr, size=7.2, color=INK, bold=(col == 2))
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    hr = h.add_run("ブラッシュアップ優先候補")
    set_font(hr, size=11, bold=True, color=BLUE)
    add_bullets(
        doc,
        [
            "P1：runtime greybox 3種類を、各テーマの完成Blenderモデルへ昇格する。",
            "P1：レバー／トグル／スロットルのpivot周辺と全可動域clearanceを視覚・実機の両方で再確認する。",
            "P1：スロットルとパワースライダーのgrip locusを造形上も明示する。",
            "P2：全8既存モデルの材質差、微細detail、目盛りをatlas中心で再設計する。",
        ],
        size=8.6,
    )
    add_shaded_note(
        doc,
        "重要",
        "緑色の配置previewは編集状態を示す一時tintであり、完成時のtheme materialではない。レビューでは通常動作時の材質とpreview表現を分けて扱う。",
        font_size=8.5,
        after=2,
    )


def add_prompt_page(doc):
    kicker = add_kicker(doc, "COPY-READY PROMPT")
    kicker.paragraph_format.page_break_before = True
    add_title(doc, "他のAIへ渡すレビュー依頼文", size=20)
    add_subtitle(
        doc,
        "この資料を添付し、以下をそのまま依頼文として使用できる。",
        size=9.5,
    )
    prompt = (
        "あなたはSF映画・アニメのメカprop、industrial design、MR操作UIに詳しい"
        "3Dアートディレクターです。添付資料にあるAnalogInstrumentMRの11種類×3テーマを"
        "レビューしてください。\n\n"
        "目的は、機能・寸法・pivot・可動範囲・Quest性能予算を維持しながら、宇宙船の"
        "ブリッジ／コックピットとして視覚的・操作的に面白い造形へ改善することです。"
        "特定作品の固有意匠は模倣しないでください。\n\n"
        "各モデルについて、(1)総合1–5点、(2)良い点、(3)最優先の変更3件、"
        "(4)silhouette／depth／mechanical plausibilityの形状案、"
        "(5)material／roughness／emission案、(6)pivot・可動干渉・grip位置の懸念、"
        "(7)3テーマ間の差とfamily統一性を回答してください。\n\n"
        "最後に、全体横断で『今すぐ直す5件』『次のart passで直す5件』『現状維持5件』、"
        "および追加で必要な正面・側面・可動端画像を列挙してください。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALL_OUT)
    p_pr.append(shd)
    for block_index, block in enumerate(prompt.split("\n\n")):
        if block_index:
            p.add_run("\n\n")
        run = p.add_run(block)
        set_font(run, size=9, color=INK)
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(2)
    hr = h.add_run("望ましい最終出力")
    set_font(hr, size=10.5, bold=True, color=BLUE)
    add_bullets(
        doc,
        [
            "11モデルの個別評価表",
            "3テーマの横断比較",
            "可動部clearanceとcontroller affordanceの指摘",
            "Quest予算内で効果が大きい改善案の優先順位",
            "次のモデル制作で必要な追加view／検証shot",
        ],
        size=8.4,
        after=2,
        line_spacing=1.1,
    )


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "AnalogInstrumentMR 3Dモデル AIレビュー資料"
    props.subject = "実装済み計器・操作部品の外部AIレビュー用資料"
    props.author = "AnalogInstrumentMR Development"
    props.keywords = "AnalogInstrumentMR, 3D model, Meta Quest, review, Unity"
    props.comments = (
        "Design preset: compact_reference_guide. "
        "Named overrides: YuGothic font, compact 9pt data tables, "
        "editorial cover, model-card page breaks."
    )


def main():
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_review_brief(doc)
    add_theme_page(doc)
    add_summary_page(doc)
    for index, model in enumerate(MODELS, 1):
        add_model_page(doc, model, index)
    add_prompt_page(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
